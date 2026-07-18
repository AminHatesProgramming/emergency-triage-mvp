from __future__ import annotations

import json
import sys
from pathlib import Path
from zipfile import BadZipFile, ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "docs" / "deliverables"
REPORT = ROOT / "reports" / "deliverable-docx-audit.json"

STALE_MARKERS = (
    "۴۸ Story/Task",
    "۷۳ رکورد",
    "AB18DA12535748BFFA8BD02D77D2B524E42154B0C9B258C9787F0E9B54EBE4D6",
    "ChatGPT",
    "Codex",
    "Notion link placeholder",
    "Original status",
)
MOJIBAKE_MARKERS = ("Ã", "Â", "Ø", "Ù", "â€", "ï¿½", "\ufffd", "提示")


def document_text(path: Path) -> str:
    document = Document(path)
    chunks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def main() -> None:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8")
    results: dict[str, dict[str, object]] = {}
    for path in sorted(DELIVERABLES.glob("*.docx")):
        item: dict[str, object] = {
            "bytes": path.stat().st_size,
            "zip_valid": False,
            "rtl_markup": False,
            "stale_markers": [],
            "mojibake_markers": [],
        }
        try:
            with ZipFile(path) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
            item["zip_valid"] = True
            item["rtl_markup"] = "<w:bidi" in document_xml or "<w:rtl" in document_xml
            text = document_text(path)
            item["stale_markers"] = [marker for marker in STALE_MARKERS if marker in text]
            item["mojibake_markers"] = [marker for marker in MOJIBAKE_MARKERS if marker in text]
        except (BadZipFile, KeyError, UnicodeDecodeError, ValueError) as exc:
            item["error"] = str(exc)
        item["passed"] = bool(
            item["zip_valid"]
            and item["rtl_markup"]
            and not item["stale_markers"]
            and not item["mojibake_markers"]
            and "error" not in item
        )
        results[path.name] = item

    required_assertions = {
        "final_report_has_current_notion_count": "۷۵ رکورد"
        in document_text(DELIVERABLES / "ITPM_Final_Report_Emergency_Triage.docx"),
        "mobile_handoff_has_current_apk_hash": "1E4FC0C64CB79EA561F9E069D37841759A5E34806734A119A1F1E350FB38FE7E"
        in document_text(DELIVERABLES / "Emdadyar_Mobile_App_For_Professor.docx"),
        "compliance_has_current_jira_count": "۵۱ Story/Task"
        in document_text(DELIVERABLES / "ITPM_Final_Announcement_Compliance_Matrix.docx"),
    }
    passed = all(item["passed"] for item in results.values()) and all(required_assertions.values())
    report = {
        "documents": results,
        "required_assertions": required_assertions,
        "visual_render_note": "LibreOffice was unavailable; structural OOXML and text QA was performed.",
        "passed": passed,
    }
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(json.dumps(report, ensure_ascii=False, indent=2))
    if not passed:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
