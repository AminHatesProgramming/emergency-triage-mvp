# -*- coding: utf-8 -*-
"""Build the final RTL Persian presentation runbook as a polished DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "final-presentation-runbook.md"
OUTPUT = ROOT / "docs" / "deliverables" / "Emdadyar_Final_Presentation_Runbook.docx"

NAVY = RGBColor(6, 78, 95)
TEAL = RGBColor(8, 127, 140)
INK = RGBColor(21, 33, 43)
MUTED = RGBColor(96, 113, 126)
LIGHT = "EAF4F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, size=11, bold=False, color=INK) -> None:
    run.font.name = "Arial"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    run._element.get_or_add_rPr().append(rtl)


def style_paragraph(paragraph, size=11, bold=False, color=INK, after=6, line=1.25) -> None:
    set_rtl(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def add_rich_text(paragraph, text: str, size=11, color=INK) -> None:
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        code = part.startswith("`") and part.endswith("`")
        clean = part[2:-2] if bold else part[1:-1] if code else part
        run = paragraph.add_run(clean)
        set_run_font(run, size=9.5 if code else size, bold=bold, color=TEAL if code else color)


def add_heading(doc, text: str, level: int) -> None:
    if level == 2 and ("بخش اول" in text or "بخش دوم" in text or "بخش سوم" in text):
        doc.add_page_break()
    p = doc.add_paragraph()
    set_rtl(p)
    sizes = {1: 24, 2: 16, 3: 12.5}
    colors = {1: NAVY, 2: NAVY, 3: TEAL}
    before = {1: 0, 2: 16, 3: 11}
    after = {1: 10, 2: 8, 3: 5}
    p.paragraph_format.space_before = Pt(before[level])
    p.paragraph_format.space_after = Pt(after[level])
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_run_font(run, size=sizes[level], bold=True, color=colors[level])


def add_table(doc, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    usable = 6.5
    if cols == 3:
        widths = [2.15, 1.05, 3.30]
    else:
        widths = [usable / cols] * cols
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.width = Inches(widths[c_idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if r_idx == 0:
                set_cell_shading(cell, LIGHT)
            p = cell.paragraphs[0]
            p.clear()
            add_rich_text(p, value, size=9.6)
            style_paragraph(p, size=9.6, bold=(r_idx == 0), color=NAVY if r_idx == 0 else INK, after=0, line=1.15)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_header_footer(doc: Document, running_title: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.text = running_title
    style_paragraph(p, size=9, bold=True, color=MUTED, after=0, line=1.0)
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "درس مدیریت پروژه فناوری اطلاعات | نسخه نهایی ۲۱ تیر ۱۴۰۵"
    style_paragraph(p, size=8.5, color=MUTED, after=0, line=1.0)


def build(source: Path = SOURCE, output: Path = OUTPUT, running_title: str = "امدادیار | سناریوی نهایی ارائه") -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    add_header_footer(doc, running_title)

    text = source.read_text(encoding="utf-8")
    text = re.sub(r"<!--.*?-->", "", text, flags=re.S)
    text = text.replace('<div dir="rtl" align="right">', "").replace("</div>", "")
    lines = text.strip().splitlines()

    table_rows: list[list[str]] = []
    in_table = False
    for raw in lines:
        line = raw.strip()
        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        if in_table:
            add_table(doc, table_rows)
            table_rows = []
            in_table = False
        if not line:
            continue
        if line.startswith("# "):
            add_heading(doc, line[2:], 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_rich_text(p, re.sub(r"^\d+\. ", "", line))
            style_paragraph(p, after=4)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_text(p, line[2:])
            style_paragraph(p, after=4)
        elif line.startswith("> "):
            p = doc.add_paragraph()
            add_rich_text(p, line[2:], size=10.5, color=NAVY)
            style_paragraph(p, size=10.5, color=NAVY, after=8)
            p.paragraph_format.left_indent = Inches(0.18)
            p.paragraph_format.right_indent = Inches(0.18)
            set_cell = OxmlElement("w:shd")
            set_cell.set(qn("w:fill"), "F1F7F7")
            p._p.get_or_add_pPr().append(set_cell)
        else:
            p = doc.add_paragraph()
            add_rich_text(p, line)
            style_paragraph(p)

    if table_rows:
        add_table(doc, table_rows)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--source", type=Path, default=SOURCE)
    parser.add_argument("--output", type=Path, default=OUTPUT)
    parser.add_argument("--header", default="امدادیار | سناریوی نهایی ارائه")
    args = parser.parse_args()
    build(args.source, args.output, args.header)
