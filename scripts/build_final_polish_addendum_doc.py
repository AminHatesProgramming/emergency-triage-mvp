from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "deliverables" / "ITPM_Final_Polish_Addendum_Emdadyar.docx"

FONT = "Tahoma"
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
FILL = "F2F4F7"
DXA_PER_INCH = 1440


def rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def font(run, size: float = 11, color: RGBColor | None = None, bold: bool = False) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_p(doc: Document, text: str, size: float = 11, bold: bool = False) -> None:
    p = doc.add_paragraph()
    rtl(p)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    font(r, size=size, bold=bold)


def add_h(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    rtl(p)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    font(r, size=15 if level == 1 else 12.5, color=BLUE if level == 1 else DARK, bold=True)


def shade(cell, fill: str = FILL) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    rtl(p)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    font(r, size=9.5, bold=bold)


def table_geometry(table, widths_in: list[float]) -> None:
    widths = [int(round(w * DXA_PER_INCH)) for w in widths_in]
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    for idx, header in enumerate(headers):
        shade(table.rows[0].cells[idx])
        cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value)


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    rtl(title)
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("پیوست روتوش نهایی پروژه امدادیار")
    font(r, size=18, color=BLUE, bold=True)

    add_p(
        doc,
        "این پیوست برای نسخه نهایی پروژه درس مدیریت پروژه فناوری اطلاعات تهیه شده و آخرین وضعیت پاکسازی مستندات، محصول، بازخورد و ابزارهای مدیریت پروژه را خلاصه می‌کند.",
    )

    add_h(doc, "خلاصه تصمیم‌های نهایی")
    add_table(
        doc,
        ["محور", "اقدام نهایی", "شاهد"],
        [
            ["مستندات", "فایل‌های draft و پوستر قدیمی حذف شدند و فهرست رسمی تحویل ساخته شد.", "docs/final-submission-index.md"],
            ["وب‌اپ", "فرم سوال‌های انتهایی حذف شد، UI روی محصول واقعی متمرکز شد و راهنمای نصب اضافه شد.", "frontend/index.html"],
            ["بازخورد", "78 بازخورد پیش‌پایلوت و 9 بازخورد پرستار تریاژ در انتظار تایید مستند شد.", "docs/triage-nurse-feedback-confirmation.md"],
            ["مدیریت پروژه", "CSVهای Jira و GitHub Project برای backlog، sprint، owner و time tracking آماده شد.", "docs/artifacts/"],
            ["مدل", "متریک‌ها بدون اغراق و با توضیح trade-off ثبت شدند.", "docs/model-final-metrics-audit.md"],
        ],
        [1.15, 3.25, 1.8],
    )

    add_h(doc, "وضعیت بازخورد پرستاران تریاژ")
    add_p(
        doc,
        "9 نظر پیشنهادی پرستار تریاژ به‌عنوان pending_confirmation ثبت شده‌اند. پس از پرسش از افراد واقعی، کد ناشناس تاییدکننده و تاریخ تایید در CSV مربوط ثبت می‌شود. این روش هم حلقه بازخورد را نشان می‌دهد و هم از جعل داده جلوگیری می‌کند.",
    )

    add_h(doc, "راهنمای ارائه به استاد")
    for item in [
        "ابتدا نسخه موبایلی امدادیار را با لینک عمومی یا local demo نمایش دهید.",
        "بعد نشان دهید پروژه فقط فنی نیست: Jira board، Sprintها، KPI، Risk و Knowledge Base دارد.",
        "در بخش بازخورد، صادقانه بگویید 78 مورد برای پیش‌بینی UX استفاده شده و 9 نظر پرستار تریاژ در مرحله تایید میدانی است.",
        "برای نقش اعضا، محمدرضا روی KPI/ریسک/بازخورد و محدثه روی UI/QA/مستندات توضیح کوتاه و قابل دفاع بدهند.",
    ]:
        add_p(doc, f"- {item}")

    add_h(doc, "چک نهایی قبل از ارسال")
    add_table(
        doc,
        ["مورد", "وضعیت"],
        [
            ["وب‌اپ امدادیار", "آماده اجرا و قابل نصب به‌صورت PWA"],
            ["گزارش Word", "موجود در پوشه deliverables"],
            ["پیوست مدیریت پروژه", "موجود و قابل ارائه"],
            ["Jira/GitHub Project", "CSV و راهنمای ورود آماده؛ ورود نهایی باید با حساب مالک انجام شود"],
            ["بازخورد پرستار تریاژ", "پیش‌نویس 9 مورد آماده؛ نیازمند تایید میدانی"],
            ["APK/AAB", "نیازمند Android Studio/SDK؛ راهنمای TWA آماده است"],
        ],
        [2.2, 4.0],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
