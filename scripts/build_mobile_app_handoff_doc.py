from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "deliverables" / "Emdadyar_Mobile_App_For_Professor.docx"
QR = ROOT / "docs" / "artifacts" / "emdadyar-pwa-qr.png"
PUBLIC_URL = "https://aminhatesprogramming.github.io/emergency-triage-mvp/"


def set_run_font(run, name: str = "Arial", size: int | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.text = ""
    set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=bold)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "D9E2EC") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_in: list[float]) -> None:
    table.autofit = False
    table.allow_autofit = False
    total_dxa = int(sum(widths_in) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    widths_dxa = [int(width * 1440) for width in widths_in]
    for width_dxa in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(widths_in[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_paragraph()
    set_paragraph_rtl(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size=14 if level == 1 else 12, bold=True)
    run.font.color.rgb = RGBColor.from_string("0B4F6C" if level == 1 else "1F4D78")
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    return paragraph


def add_body(doc: Document, text: str, *, bold: bool = False):
    paragraph = doc.add_paragraph()
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.2
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=bold)
    return paragraph


def add_hyperlink(paragraph, url: str, text: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10)

    title = doc.add_paragraph()
    set_paragraph_rtl(title, WD_ALIGN_PARAGRAPH.CENTER)
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("نسخه موبایلی قابل نصب «امدادیار»")
    set_run_font(run, size=20, bold=True)
    run.font.color.rgb = RGBColor.from_string("063B4A")

    subtitle = doc.add_paragraph()
    set_paragraph_rtl(subtitle, WD_ALIGN_PARAGRAPH.CENTER)
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run("سامانه پشتیبان تصمیم‌گیری تریاژ اورژانس - پروژه درس مدیریت پروژه فناوری اطلاعات")
    set_run_font(run, size=10)
    run.font.color.rgb = RGBColor.from_string("52616B")

    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(top_table, [4.45, 2.05])
    set_table_borders(top_table, "BFD7E3")
    left, right = top_table.rows[0].cells
    shade_cell(left, "F4FAFC")
    shade_cell(right, "FFFFFF")

    set_cell_text(left, "لینک مستقیم نسخه عمومی:", bold=True, color="063B4A")
    link_p = left.add_paragraph()
    set_paragraph_rtl(link_p)
    add_hyperlink(link_p, PUBLIC_URL, PUBLIC_URL)
    note_p = left.add_paragraph()
    set_paragraph_rtl(note_p)
    note_run = note_p.add_run("این نسخه بدون سرور local روی موبایل و دسکتاپ باز می‌شود و پس از اولین باز شدن، مانند وب‌اپ قابل نصب است.")
    set_run_font(note_run, size=9)
    note_run.font.color.rgb = RGBColor.from_string("334E68")

    if QR.exists():
        qr_p = right.paragraphs[0]
        qr_p.text = ""
        qr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        qr_p.add_run().add_picture(str(QR), width=Inches(1.6))
        cap = right.add_paragraph()
        set_paragraph_rtl(cap, WD_ALIGN_PARAGRAPH.CENTER)
        cap_run = cap.add_run("QR Code لینک وب‌اپ")
        set_run_font(cap_run, size=9, bold=True)

    add_heading(doc, "روش نصب روی گوشی اندروید")
    for item in [
        "لینک بالا را با Chrome روی گوشی باز کنید.",
        "از منوی سه‌نقطه گزینه Install app یا Add to Home screen را انتخاب کنید.",
        "آیکن «امدادیار» روی صفحه اصلی گوشی اضافه می‌شود و مثل یک اپلیکیشن اجرا می‌شود.",
    ]:
        add_body(doc, f"• {item}")

    add_heading(doc, "وضعیت نسخه قابل ارسال")
    status_rows = [
        ("اجرا بدون سرور local", "آماده"),
        ("نمایش روی موبایل", "آماده"),
        ("نصب به‌صورت PWA", "آماده"),
        ("نیاز به حساب کاربری", "ندارد"),
        ("نیاز به بک‌اند آنلاین", "ندارد؛ مدل سبک در مرورگر اجرا می‌شود"),
        ("مناسب برای جمع‌آوری بازخورد پایلوت", "بله"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    headers = table.rows[0].cells
    set_cell_text(headers[0], "بخش", bold=True, color="FFFFFF")
    set_cell_text(headers[1], "وضعیت", bold=True, color="FFFFFF")
    shade_cell(headers[0], "0B4F6C")
    shade_cell(headers[1], "0B4F6C")
    for key, value in status_rows:
        row = table.add_row().cells
        set_cell_text(row[0], key)
        set_cell_text(row[1], value, bold=value in {"آماده", "بله"})
        shade_cell(row[0], "F7FBFC")
        shade_cell(row[1], "FFFFFF")
    set_table_geometry(table, [3.1, 3.4])

    add_heading(doc, "نکته اخلاقی و آموزشی")
    add_body(
        doc,
        "امدادیار یک نمونه آموزشی از سامانه پشتیبان تصمیم‌گیری تریاژ است. خروجی سامانه جایگزین پزشک، پرستار یا پروتکل رسمی درمانی نیست و فقط برای نمایش مسیر محصول، ارزیابی اولیه و تصمیم‌سازی در پروژه درس مدیریت پروژه فناوری اطلاعات ارائه شده است.",
    )

    add_heading(doc, "سناریوی پیشنهادی برای تست")
    add_body(
        doc,
        "برای نمایش خروجی، چند داده حیاتی بیمار را وارد کنید و دکمه «ارزیابی بیمار» را بزنید. حتی با چند ورودی محدود نیز سامانه یک ارزیابی اولیه ارائه می‌دهد و میزان کامل بودن اطلاعات را جداگانه نمایش می‌دهد.",
    )

    footer = section.footer.paragraphs[0]
    set_paragraph_rtl(footer, WD_ALIGN_PARAGRAPH.CENTER)
    footer_run = footer.add_run("Emdadyar PWA | Public mobile web app | Decision support only")
    set_run_font(footer_run, size=8)
    footer_run.font.color.rgb = RGBColor.from_string("7B8794")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_docx()
