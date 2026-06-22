from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "deliverables" / "Emdadyar_Mobile_App_For_Professor.docx"
PUBLIC_URL = "https://aminhatesprogramming.github.io/emergency-triage-mvp/"


def main() -> None:
    doc = Document(DOCX)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    with ZipFile(DOCX) as archive:
        names = archive.namelist()
        document_xml = archive.read("word/document.xml").decode("utf-8")
        rels_xml = archive.read("word/_rels/document.xml.rels").decode("utf-8")

    checks = {
        "exists": DOCX.exists(),
        "size_gt_10kb": DOCX.stat().st_size > 10_000,
        "paragraphs": len(doc.paragraphs) >= 10,
        "tables": len(doc.tables) >= 2,
        "has_title": "نسخه موبایلی قابل نصب" in text,
        "has_public_url_text": "aminhatesprogramming.github.io/emergency-triage-mvp" in document_xml,
        "has_public_url_rel": PUBLIC_URL in rels_xml,
        "has_qr_image": any(name.startswith("word/media/") for name in names),
        "has_rtl_bidi": "<w:bidi" in document_xml,
        "has_table_widths": "<w:tblW" in document_xml and "<w:gridCol" in document_xml,
    }

    for key, value in checks.items():
        print(f"{key}: {value}")

    failed = [key for key, value in checks.items() if not value]
    if failed:
        raise SystemExit(f"Failed checks: {', '.join(failed)}")


if __name__ == "__main__":
    main()
