from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DELIVERABLES = ROOT / "docs" / "deliverables"
ACCENT = RGBColor(0x0B, 0x67, 0x70)


def make_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))
    for run in paragraph.runs:
        run.font.name = "Arial"
        r_pr = run._r.get_or_add_rPr()
        if r_pr.find(qn("w:rtl")) is None:
            r_pr.append(OxmlElement("w:rtl"))


def replace_paragraph(document: Document, marker: str, text: str) -> None:
    if any(paragraph.text == text for paragraph in document.paragraphs):
        return
    for paragraph in document.paragraphs:
        if marker in paragraph.text:
            paragraph.clear()
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(11)
            make_rtl(paragraph)
            return
    raise RuntimeError(f"Paragraph marker not found: {marker}")


def add_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT
    make_rtl(paragraph)


def add_body(document: Document, text: str, *, bullet: bool = False) -> None:
    style = "List Bullet" if bullet else None
    paragraph = document.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    make_rtl(paragraph)


def update_final_report() -> None:
    path = DELIVERABLES / "ITPM_Final_Report_Emergency_Triage.docx"
    document = Document(path)
    replace_paragraph(
        document,
        "backlog آماده Jira شامل ۹ Epic",
        "پروژه واقعی Jira با کلید EMD شامل ۹ Epic و ۵۱ Story/Task است؛ همه Story/Taskها به Epic والد متصل شده‌اند و وضعیت Workflow بر اساس شواهد repo به‌روزرسانی شده است.",
    )
    replace_paragraph(
        document,
        "انتقال backlog قدیمی به Jira",
        "ساعت‌های ۴۷، ۲۳ و ۱۸، برآورد مبتنی بر فعالیت‌های ثبت‌شده‌اند و ادعای تایم‌شیت لحظه‌ای از ابتدای ترم نیستند. backlog قدیمی در Jira با اتکا به شواهد Git و فایل‌ها بازسازی شده است؛ تاریخ ایجاد پلتفرم به‌عنوان تاریخ انجام کار معرفی نمی‌شود. حساب Jira محدثه و محمدرضا برای ثبت Assignee و worklog شخصی هنوز باید با ایمیل واقعی دعوت شود.",
    )
    replace_paragraph(
        document,
        "لینک و screenshot واقعی Jira/Notion",
        "ابزارهای واقعی فعال‌اند: Jira در https://pourmand.atlassian.net/jira/software/projects/EMD/board و Notion در https://app.notion.com/p/38fd955c965a80c18b7ac3a8fd176cc3. ثبت screenshot رابط، آزمون دسترسی مهمان و دعوت حساب دو عضو همچنان اقدام انسانی پیش از تحویل است.",
    )
    marker = "ضمیمه ممیزی فنی نسخه انتشار"
    if not any(marker in paragraph.text for paragraph in document.paragraphs):
        document.add_page_break()
        add_heading(document, marker)
        add_body(
            document,
            "در ۲۵ تیر ۱۴۰۵، پیش از بسته‌بندی نسخه بازار، ارزیابی مدل و قواعد ایمنی از ابتدا بازتولید شد. "
            "این ممیزی داخلی و گذشته‌نگر است و جای اعتبارسنجی بالینی یا تطبیق با پروتکل بیمارستان را نمی‌گیرد.",
        )
        evidence = [
            "تقسیم آزمون با random_state=42 بازسازی و ۱۱۱٬۶۰۶ رکورد نگه‌داشته‌شده دوباره پیش‌بینی شد؛ اختلاف متریک با artifact آموزشی صفر بود.",
            "مجموعه QA شامل ۱۱۸ تست Python، ۱۰۷ آزمون مرزی سن‌محور و ۵٬۰۰۰ ترکیب تصادفی با بذر ثابت است.",
            "نسخه مرورگر و API در ۱٬۱۶۷ سناریو مقایسه شدند؛ صفر اختلاف تصمیمی و بیشینه اختلاف احتمال 3.54e-7 ثبت شد.",
            "برای ورودی سه‌فیلدی شکایت اصلی، ضربان و SpO2، Recall برابر 0.9411 و FPR برابر 0.5901 بود؛ افزودن سن Recall=0.9246 و FPR=0.5177 داد.",
            "آستانه آزمایشی فقط علائم حیاتی به علت FPR=0.7840 رد شد و وارد محصول نشد.",
            "قواعد ایمنی نسخه 2026.07.2 سن، اکسیژن، فشار، تنفس، ضربان، دما، نسبت شوک و ناسازگاری اندازه‌گیری را بررسی می‌کنند؛ احتمال خام مدل تغییر نمی‌کند.",
            "APK و AAB نسخه 1.0.0 با assetهای نهایی بازامضا شدند؛ پنج فایل کلیدی داخل بسته بایت‌به‌بایت با dist برابرند.",
            "نصب روی گوشی فیزیکی، تحلیل fairness و اعتبارسنجی خارجی/بالینی همچنان اقدام‌های باز فاز بعدی هستند.",
        ]
        for item in evidence:
            add_body(document, item, bullet=True)
        add_body(
            document,
            "شواهد: docs/model-release-scenario-audit-fa.md، reports/model/release_validation_v7.json، "
            "reports/model/browser_backend_differential_v7.json و docs/market/build-verification-fa.md.",
        )
    pm_marker = "شواهد ابزارهای واقعی مدیریت پروژه و دانش"
    if not any(pm_marker in paragraph.text for paragraph in document.paragraphs):
        add_heading(document, pm_marker)
        items = [
            "Jira: پروژه EMD با ۹ Epic، ۵۱ Story/Task، Parent واقعی و ۵۲ مورد Done؛ کارهای بیرونی اثبات‌نشده باز مانده‌اند.",
            "Notion: هفت صفحه محتوایی، پنج دیتابیس و ۶۶ رکورد تصمیم، تغییر، ریسک، بازخورد و QA.",
            "بازخورد: ۹ نظر کیفی پرستاران تأیید شده؛ outreach دو مرکز درمانی در انتظار پاسخ نهایی است و همکاری رسمی ادعا نمی‌شود.",
            "کارهای باز: حساب Jira دو عضو، screenshot، تست APK/PWA روی گوشی، ضبط ویدئو، fairness و اعتبارسنجی بالینی.",
        ]
        for item in items:
            add_body(document, item, bullet=True)
    document.save(path)


def update_runbook() -> None:
    path = DELIVERABLES / "Emdadyar_Final_Presentation_Runbook.docx"
    document = Document(path)
    replace_paragraph(
        document,
        "Jira/Notion package",
        "اپ عمومی، نمودار مدل، Burndown، برد واقعی Jira، دانشنامه واقعی Notion و گزارش بازخورد را در تب‌های جدا باز کنید.",
    )
    replace_paragraph(
        document,
        "مدل نهایی v7 فقط از اطلاعات زمان تریاژ استفاده می‌کند",
        "«مدل نهایی v7 فقط از اطلاعات زمان تریاژ استفاده می‌کند؛ داده‌های بعدی مانند آزمایش، تصویربرداری، دارو و نتیجه بستری حذف شده‌اند تا data leakage رخ ندهد. Threshold روی validation انتخاب شده و بازتولید روی ۱۱۱٬۶۰۶ رکورد test همان AUC برابر ۰٫۹۰۴۱ و Recall برابر ۰٫۹۲۴۶ را بدون اختلاف تأیید کرد. قواعد ایمنی احتمال مدل را تغییر نمی‌دهند و فوریت را جداگانه ثبت می‌کنند.»",
    )
    replace_paragraph(
        document,
        "تست‌ها شامل فرم خالی، بیمار پرخطر",
        "«QA فقط چند سناریوی نمایشی نبود: ۱۱۸ تست Python، ۱۰۷ مرز سن‌محور، ۵٬۰۰۰ ترکیب تصادفی و ۱٬۱۶۷ مقایسه مرورگر/API اجرا شد و اختلاف تصمیمی صفر بود. سناریوهای فرم خالی، بیمار پرخطر، وضعیت متوسط، داده کم، درد قفسه سینه و پاک‌کردن فرم نیز در Test Log ثبت شده‌اند. نصب APK روی گوشی فیزیکی را صادقانه به‌عنوان کار باز نگه داشته‌ایم.»",
    )
    replace_paragraph(
        document,
        "چرا FPR برابر ۰٫۳۳۵۲ است؟",
        "چرا FPR برابر ۰٫۳۳۵۲ است؟ حالت پیش‌فرض برای حفظ Recall بالا انتخاب شده است. برای ورودی ناقص FPR بالاتر نیز جدا گزارش شده و گزینه فقط علائم حیاتی با FPR برابر ۰٫۷۸۴۰ عمداً از محصول کنار گذاشته شد؛ تنظیم نهایی باید با ظرفیت و پروتکل مرکز درمانی انجام شود.",
    )
    links_marker = "لینک‌های آماده پیش از ضبط"
    if not any(links_marker in paragraph.text for paragraph in document.paragraphs):
        add_heading(document, links_marker)
        links = [
            "وب‌اپ: https://aminhatesprogramming.github.io/emergency-triage-mvp/",
            "Jira: https://pourmand.atlassian.net/jira/software/projects/EMD/board",
            "Notion: https://app.notion.com/p/38fd955c965a80c18b7ac3a8fd176cc3",
            "GitHub: https://github.com/AminHatesProgramming/emergency-triage-mvp",
        ]
        for item in links:
            add_body(document, item, bullet=True)
    document.save(path)


def update_pm_evidence() -> None:
    path = DELIVERABLES / "ITPM_Project_Management_Evidence_Package.docx"
    document = Document(path)
    replace_paragraph(
        document,
        "برای نمره کامل، همین ساختار باید",
        "ساختار در ابزارهای واقعی اجرا شده است: پروژه EMD در Jira برای Work Management و صفحه «امداد یار | دانشنامه پروژه و مدیریت دانش» در Notion برای Knowledge Management. Storyها به Epic والد، شواهد Git و رکوردهای دانش متصل شده‌اند.",
    )
    replace_paragraph(
        document,
        "ورود/اسکرین‌شات نهایی Jira و Notion",
        "ریسک باز: اعتبارسنجی بالینی، fairness، نصب دستگاه واقعی، screenshot رابط Jira/Notion، آزمون دسترسی مهمان و دعوت حساب Jira دو عضو تیم.",
    )
    marker = "کنترل انتشار و شواهد نهایی"
    if not any(marker in paragraph.text for paragraph in document.paragraphs):
        document.add_page_break()
        add_heading(document, marker)
        items = [
            "Deliverable فنی: مدل v7، وب‌اپ عمومی، APK/AAB امضاشده و بسته بازار نسخه 1.0.0.",
            "Evidence ارزیابی: ۱۱۱٬۶۰۶ رکورد test، ۱۱۸ تست Python و ۱٬۱۶۷ سناریوی تطابق مرورگر/API.",
            "تصمیم کنترل دامنه: آستانه core-vitals با FPR=0.7840 رد شد؛ فقط دو الگوی sparse دارای شواهد validation فعال‌اند.",
            "Evidence ذی‌نفع: ۹ بازخورد کیفی تأییدشده پرستاران؛ دو درخواست ملاقات درمانی در انتظار پاسخ نهایی.",
            "ریسک باز: اعتبارسنجی بالینی، fairness، نصب دستگاه واقعی، screenshot رابط و حساب Jira دو عضو تیم.",
            "مالک شواهد فنی: محمدامین؛ مالک QA و مستندسازی: محدثه؛ مالک KPI، ریسک و پیگیری ذی‌نفع: محمدرضا.",
        ]
        for item in items:
            add_body(document, item, bullet=True)
    tools_marker = "وضعیت ابزارهای واقعی"
    if not any(tools_marker in paragraph.text for paragraph in document.paragraphs):
        add_heading(document, tools_marker)
        items = [
            "Jira Project: https://pourmand.atlassian.net/jira/software/projects/EMD/board؛ ۹ Epic و ۵۱ Story/Task.",
            "Notion Home: https://app.notion.com/p/38fd955c965a80c18b7ac3a8fd176cc3؛ ۷ صفحه، ۵ دیتابیس و ۶۶ رکورد.",
            "۵۲ Issue در وضعیت Done، پنج مورد In Progress، یک مورد In Review و دو Task بیرونی در Backlog ثبت شده‌اند.",
            "Labelهای owner برای هر سه عضو ثبت شده‌اند؛ Assignee مستقیم محدثه و محمدرضا پس از دعوت حساب واقعی تکمیل می‌شود.",
        ]
        for item in items:
            add_body(document, item, bullet=True)
    document.save(path)


def update_mobile_handoff() -> None:
    path = DELIVERABLES / "Emdadyar_Mobile_App_For_Professor.docx"
    document = Document(path)
    marker = "فایل نصب مستقیم Android"
    if not any(marker in paragraph.text for paragraph in document.paragraphs):
        add_heading(document, marker)
        add_body(
            document,
            "نسخه امضاشده APK در بسته تحویل با نام Emdadyar-1.0.0-release.apk قرار دارد. "
            "این نسخه مدل و رابط را داخل برنامه نگه می‌دارد و برای ارزیابی بیمار به اینترنت یا سرور محلی وابسته نیست.",
        )
        details = [
            "شناسه بسته: ir.pourmand.emdadyar؛ نسخه 1.0.0؛ حداقل Android 6 (API 23).",
            "SHA-256 فایل APK: 136F57F1CDA715AE53C37E0A279F8A3ABB43752098FC5321C99F8829A540C93F.",
            "فایل AAB برای پنل مارکت است و مستقیماً روی گوشی نصب نمی‌شود.",
            "برای نصب مستقیم APK ممکن است لازم باشد اجازه نصب از منبع انتخاب‌شده در تنظیمات Android فعال شود.",
            "امضا و دارایی‌های داخلی تأیید شده‌اند؛ نصب روی یک گوشی فیزیکی همچنان باید پیش از انتشار عمومی ثبت شود.",
        ]
        for item in details:
            add_body(document, item, bullet=True)
    document.save(path)


if __name__ == "__main__":
    update_final_report()
    update_runbook()
    update_pm_evidence()
    update_mobile_handoff()
    print("Updated final report, presentation runbook, PM evidence, and mobile handoff.")
