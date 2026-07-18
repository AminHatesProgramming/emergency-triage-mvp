from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "deliverables"
OUT_PATH = OUT_DIR / "ITPM_Group_Project_Info_Emdadyar.docx"

FONT = "Tahoma"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
DXA_PER_INCH = 1440

PROJECT_DESCRIPTION = (
    "امداد یار یک وب‌اپ پشتیبان تصمیم برای ارزیابی اولیه بیماران اورژانسی است. "
    "سامانه با دریافت علائم حیاتی، شکایت اصلی و سوابق قابل دسترس، احتمال نیاز به بررسی فوری را تخمین می‌زند. "
    "خروجی فقط برای کمک به تصمیم کادر درمان است و جایگزین پزشک یا پرستار نیست."
)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(10.5)


def set_table_geometry(table, column_widths_in: list[float]) -> None:
    widths_dxa = [int(round(width * DXA_PER_INCH)) for width in column_widths_in]
    table_width_dxa = sum(widths_dxa)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(table_width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_grid = table._tbl.tblGrid
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")


def set_paragraph_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_run_font(run, size: float = 11, color: RGBColor | None = None, bold: bool = False) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=13, color=BLUE, bold=True)


def add_body_paragraph(doc: Document, text: str, after: int = 6) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_rtl(paragraph)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_run_font(run, size=11)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_table_geometry(table, widths)

    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.width = Inches(widths[idx])
        set_cell_shading(cell, GRAY_FILL)
        set_cell_text(cell, header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].width = Inches(widths[idx])
            set_cell_text(cells[idx], value)

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.line_spacing = 1.15

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def build() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    set_paragraph_rtl(title)
    title.paragraph_format.space_after = Pt(3)
    title_run = title.add_run("اطلاعات تکمیلی گروه پروژه")
    set_run_font(title_run, size=18, color=DARK_BLUE, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_rtl(subtitle)
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("درس مدیریت پروژه فناوری اطلاعات | پروژه امداد یار")
    set_run_font(subtitle_run, size=11, color=RGBColor(85, 85, 85))

    add_heading(doc, "۱. مشخصات تیم")
    add_table(
        doc,
        headers=["نام و نام خانوادگی", "نقش در پروژه", "توضیح"],
        rows=[
            ["محمدامین پورمند", "مدیر پروژه / معمار سیستم و مدل", "شماره دانشجویی مدیر پروژه: [وارد شود]"],
            ["محمدرضا آرمان پور", "هماهنگ‌کننده کنترل پروژه و شاخص‌ها", "عضو تیم"],
            ["محدثه حاتمی کیا", "هماهنگ‌کننده مستندات، رابط کاربری و QA", "عضو تیم"],
        ],
        widths=[2.0, 2.35, 2.15],
    )

    add_heading(doc, "۲. اطلاعات تماس")
    add_table(
        doc,
        headers=["نوع اطلاعات", "مقدار"],
        rows=[
            ["ایمیل معتبر", "[وارد شود]"],
            ["شماره تماس", "[وارد شود]"],
        ],
        widths=[2.0, 4.5],
    )

    add_heading(doc, "۳. شرح پروژه")
    add_body_paragraph(doc, PROJECT_DESCRIPTION, after=6)
    add_body_paragraph(doc, "تعداد کلمات شرح پروژه: ۴۵ کلمه", after=0)

    doc.core_properties.title = "اطلاعات تکمیلی گروه پروژه امداد یار"
    doc.core_properties.subject = "درس مدیریت پروژه فناوری اطلاعات"
    doc.core_properties.author = "محمدامین پورمند"
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
