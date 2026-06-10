from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "deliverables"
OUT.mkdir(parents=True, exist_ok=True)

FONT = "IRANSans(FaNum)"
FALLBACK_FONT = "Tahoma"
TEAL = "087F8C"
NAVY = "0B2545"
LIGHT_TEAL = "EAF7F7"
LIGHT_GRAY = "F3F6F8"
RED = "C93535"
GREEN = "17845B"
AMBER = "B77A0B"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = "D9E4E8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_paragraph_rtl(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def style_doc(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    for name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), FONT)

    normal = styles["Normal"]
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    h1 = styles["Heading 1"]
    h1.font.size = Pt(17)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(TEAL)
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(7)

    h2 = styles["Heading 2"]
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(NAVY)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "Emergency Triage Decision Support | ITPM Final Package"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        set_run_font(run, 8, False, "667781")

    footer = section.footer.paragraphs[0]
    footer.text = "Decision-support only; not a replacement for clinical judgment."
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        set_run_font(run, 8, False, "667781")

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(title)
    set_run_font(run, 22, True, NAVY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    set_run_font(run, 12, False, TEAL)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("درس مدیریت و کنترل پروژه‌های فناوری اطلاعات | نسخه رسمی برای تحویل")
    set_run_font(run, 10, False, "667781")

    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    set_paragraph_rtl(p)


def add_para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_rtl(p)
    if bold_lead:
        run = p.add_run(bold_lead)
        set_run_font(run, 10.5, True, NAVY)
        run = p.add_run(text)
        set_run_font(run, 10.5)
    else:
        run = p.add_run(text)
        set_run_font(run, 10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_rtl(p)
        run = p.add_run(item)
        set_run_font(run, 10.2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = ""
        set_cell_shading(hdr[i], TEAL)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, 9.5, True, "FFFFFF")
        if widths:
            set_cell_width(hdr[i], widths[i])
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2 == 0:
                set_cell_shading(cells[i], "FBFCFD")
            p = cells[i].paragraphs[0]
            set_paragraph_rtl(p)
            if i == len(row) - 1 and text in {"سبز", "Pass", "آماده"}:
                color = GREEN
            elif text in {"زرد", "نیازمند polish نهایی"}:
                color = AMBER
            elif text in {"قرمز"}:
                color = RED
            else:
                color = "17212B"
            r = p.add_run(str(text))
            set_run_font(r, 8.7, False, color)
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph()


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="BBDADD")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_rtl(p)
    r = p.add_run(title + ": ")
    set_run_font(r, 10.2, True, NAVY)
    r = p.add_run(body)
    set_run_font(r, 10.2)
    doc.add_paragraph()


def add_picture_if_exists(doc: Document, path: Path, width_cm: float, caption: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, 8.5, False, "667781")


def final_report() -> None:
    doc = Document()
    style_doc(
        doc,
        "گزارش نهایی پروژه: سیستم هوشمند تریاژ اورژانس",
        "MVP تصمیم‌یار، API-first، mobile-first و safety-first",
    )

    add_heading(doc, "خلاصه مدیریتی")
    add_para(
        doc,
        "این پروژه یک MVP مبتنی بر هوش مصنوعی برای کمک به تریاژ اورژانس است. سامانه با دریافت داده‌های قابل دسترسی در لحظه تریاژ، احتمال بحرانی بودن بیمار را تخمین می‌زند و همراه با سطح اعتماد و توضیح کوتاه ارائه می‌کند.",
    )
    add_callout(
        doc,
        "اصل اخلاقی",
        "سامانه فقط decision-support است و جایگزین پزشک، پرستار یا پروتکل رسمی بیمارستان نیست.",
    )

    add_heading(doc, "اعضای تیم و نقش‌ها")
    add_table(
        doc,
        ["عضو", "نقش", "مسئولیت قابل ارائه"],
        [
            ["محمدامین پورمند", "Project Lead / ML & System Architect", "مدل، API، معماری، GitHub و یکپارچه‌سازی"],
            ["محمدرضا آرمان پور", "Project Control & Metrics Coordinator", "KPI، ریسک، ارزش اجتماعی و trade-offها"],
            ["محدثه حاتمی کیا", "UI/Documentation & QA Coordinator", "user flow، QA، مستندات و اخلاق AI"],
        ],
        [4.2, 5.0, 7.2],
    )

    add_heading(doc, "Scope و تغییر موضوع")
    add_para(
        doc,
        "در Sheet1 اولیه، موضوع گروه روی پلتفرم ارتباطات اجتماعی از طریق بازی‌های گروهی ثبت شده بود. تیم در ادامه با توجه به تاکید درس بر بحران، اثر اجتماعی و MVP قابل دفاع، scope را به سیستم هوشمند تریاژ اورژانس تغییر داد.",
    )
    add_bullets(
        doc,
        [
            "اثر انسانی و ملی مستقیم‌تر در شرایط بحران درمانی",
            "امکان تعریف KPIهای کمی روشن مثل AUC، Recall و FPR",
            "امکان ساخت demo واقعی با API و UI",
            "امکان مستندسازی قوی‌تر برای ریسک، اخلاق و مدیریت پروژه",
        ],
    )

    add_heading(doc, "محصول ساخته‌شده")
    add_bullets(
        doc,
        [
            "مدل ML نسخه v6 با threshold safety-first",
            "Backend با FastAPI شامل endpointهای /health، /model-info و /predict",
            "Frontend فارسی و mobile-first برای demo",
            "نسخه PWA قابل نصب روی موبایل با manifest، service worker و app shell آفلاین",
            "لایه safety-first hybrid برای تشخیص red flagهای حیاتی حتی در ورودی کم‌داده",
            "پشتیبانی از ورودی ناقص همراه با data completeness و missing fields",
            "مستندات کامل مدیریت پروژه، AI usage، Model Card و گزارش نهایی",
        ],
    )

    add_heading(doc, "نسخه موبایل و PWA")
    add_bullets(
        doc,
        [
            "UI برای صفحه کوچک بازطراحی شد: action bar پایین صفحه، تب‌های ورودی/خروجی/پروژه و پنل نتیجه قابل اسکن.",
            "manifest.webmanifest و service worker اضافه شد تا اپ بعد از اولین بارگذاری، app shell را در حالت آفلاین نیز باز کند.",
            "سناریوهای demo شامل بحرانی، متوسط، کم‌داده و قلبی شدند تا در ویدئو و ارائه، رفتار سیستم با داده ناقص قابل نمایش باشد.",
            "در خود MVP لینک مستقیم به بسته شواهد مدیریت پروژه اضافه شد تا محصول، مستندات و Rubric به هم متصل باشند.",
        ],
    )

    add_heading(doc, "لایه Safety-first Hybrid")
    add_table(
        doc,
        ["مولفه خروجی", "توضیح مدیریتی/فنی"],
        [
            ["model_probability", "احتمال خام مدل v6 بدون دستکاری"],
            ["critical_probability", "احتمال عملیاتی پس از اعمال guardهای ایمنی شفاف"],
            ["safety_flags", "هشدارهای بالینی ساده مثل SpO2 پایین، فشار بسیار پایین یا درد قفسه سینه با سابقه قلبی"],
            ["next_best_actions", "اقدام بعدی قابل ارائه به پرسنل؛ تصمیم نهایی همچنان انسانی است"],
        ],
        [5.0, 10.5],
    )

    add_heading(doc, "داده و کنترل Leakage")
    add_para(
        doc,
        "مدل فقط از داده‌های triage-time یا EHR-safe استفاده می‌کند: علائم حیاتی اولیه، شکایت اصلی، سن، روش ورود، سوابق مراجعه/بستری/جراحی و سابقه‌های بالینی قابل پرسش.",
    )
    add_para(
        doc,
        "داده‌های بعد از تریاژ مثل آزمایش، دارو، تصویربرداری، disposition و تشخیص‌های بعدی حذف شدند. همچنین race، ethnicity و insurance تا قبل از fairness review وارد مدل نشده‌اند.",
    )

    add_heading(doc, "نتایج مدل v6")
    add_table(
        doc,
        ["KPI", "مقدار تست", "تفسیر"],
        [
            ["AUC", "0.8947", "کیفیت کلی رتبه‌بندی مدل"],
            ["Average Precision", "0.8034", "مناسب برای کلاس بحرانی"],
            ["Recall", "0.9241", "عبور از هدف 0.92"],
            ["Precision", "0.5269", "trade-off حالت safety-first"],
            ["FPR", "0.3598", "کمتر از نسخه‌های اولیه با FPR حدود 0.487"],
            ["Threshold", "0.2962", "انتخاب‌شده روی validation"],
        ],
        [4.0, 3.0, 9.0],
    )
    add_picture_if_exists(doc, ROOT / "reports" / "model" / "confusion_matrix_v6.png", 12.5, "Confusion Matrix مدل v6")
    add_picture_if_exists(doc, ROOT / "reports" / "model" / "confidence_distribution_v6.png", 14.0, "توزیع confidence در split تست")

    add_heading(doc, "رویکرد Agile و مدیریت پروژه")
    add_table(
        doc,
        ["Sprint", "هدف", "خروجی"],
        [
            ["Sprint 0", "انتخاب مسئله و scope", "problem statement و ارزش اجتماعی"],
            ["Sprint 1", "مدل و مستندات پایه", "baseline و docs اولیه"],
            ["Sprint 2", "API و UI demo", "FastAPI و UI فارسی"],
            ["Sprint 3", "مدل v6 و مستندات", "metrics نهایی و Model Card"],
            ["Sprint 4", "پوستر و تحویل نهایی", "Word/PDF، ویدئو، board و knowledge base"],
        ],
        [3.0, 5.5, 7.5],
    )
    add_table(
        doc,
        ["شاخص مدیریت پروژه", "مقدار"],
        [
            ["مجموع زمان ثبت‌شده", "88 ساعت"],
            ["Story point اولیه", "42"],
            ["Story point باقی‌مانده", "2"],
            ["Velocity میانگین", "حدود 8 story point"],
            ["تعداد اسناد اصلی", "29 سند"],
        ],
        [7.0, 7.0],
    )

    add_heading(doc, "ریسک‌ها و کنترل‌ها")
    add_table(
        doc,
        ["ریسک", "کنترل"],
        [
            ["False Negative در بیمار بحرانی", "threshold safety-first و تمرکز روی Recall"],
            ["اتکای بیش از حد به AI", "disclaimer و تاکید بر decision-support"],
            ["Leakage", "حذف داده‌های بعد از تریاژ"],
            ["Bias جمعیتی", "حذف متغیرهای حساس تا قبل از fairness review"],
            ["اعتبار بالینی", "ثبت محدودیت و نیاز به بازخورد متخصص"],
        ],
        [7.0, 9.0],
    )

    add_heading(doc, "Roadmap آینده")
    add_table(
        doc,
        ["فاز", "هدف", "KPI"],
        [
            ["تا 12 تیر", "تحویل نهایی درس", "پوستر، ویدئو، گزارش و board"],
            ["تا پایان تیر", "اعتبارسنجی دانشگاهی", "35 کاربر آزمایشی و 2 بازخورد متخصص"],
            ["مرداد", "اعتبارسنجی فنی/اخلاقی", "False Negative و fairness analysis"],
            ["شهریور", "اعتبارسنجی نسخه PWA و local-first", "تست میدانی با فرم موبایل و بازخورد متخصص"],
        ],
        [3.2, 7.0, 5.8],
    )

    add_heading(doc, "جمع‌بندی")
    add_para(
        doc,
        "پروژه علاوه بر مدل ML، یک بسته کامل قابل ارائه برای درس مدیریت پروژه فناوری اطلاعات دارد: محصول قابل demo، متریک‌های کمی، مستندات Agile، ریسک، KPI، مدیریت دانش و برنامه آینده.",
    )

    doc.save(OUT / "ITPM_Final_Report_Emergency_Triage.docx")


def management_package() -> None:
    doc = Document()
    style_doc(
        doc,
        "پیوست مدیریت پروژه و شواهد نمره‌دهی",
        "KPI، Burndown، RACI، WBS، Time Tracking، Risk و Knowledge Base",
    )

    add_heading(doc, "چک‌لیست تطبیق با Rubric")
    add_table(
        doc,
        ["الزام استاد", "خروجی پروژه", "وضعیت"],
        [
            ["پوستر A0", "docs/poster/current-poster.png", "آماده"],
            ["تصاویر UI", "UI موبایل/PWA و سناریوهای demo", "آماده"],
            ["نسخه موبایل", "manifest، service worker، app icon و action bar موبایل", "آماده"],
            ["ماتریس همکاری و ساعات", "team-collaboration-matrix و این پیوست", "آماده"],
            ["Burndown", "docs/artifacts/burndown.svg و CSV", "آماده"],
            ["KPI فعلی و آینده", "kpi-register", "آماده"],
            ["User Acquisition Forecast", "user-acquisition.svg و CSV", "آماده"],
            ["Work Management", "work-management-board", "آماده"],
            ["Knowledge Management", "knowledge-management-index", "آماده"],
            ["ویدئوی 10 دقیقه‌ای", "video-presentation-script", "آماده برای ضبط"],
            ["گزارش نهایی", "DOCX رسمی", "آماده"],
        ],
        [6.0, 7.0, 3.0],
    )

    add_heading(doc, "KPI Register")
    add_table(
        doc,
        ["KPI", "مقدار", "هدف", "وضعیت"],
        [
            ["AUC", "0.8947", ">= 0.87", "سبز"],
            ["Average Precision", "0.8034", ">= 0.78", "سبز"],
            ["Recall بحرانی", "0.9241", ">= 0.92", "سبز"],
            ["FPR", "0.3598", "< 0.487", "سبز"],
            ["API health", "Pass", "Pass", "سبز"],
            ["ورودی ناقص", "Pass", "3-4 آیتم", "سبز"],
            ["PWA shell", "Pass", "static/offline shell", "سبز"],
            ["Safety flags", "Pass", "red flag output", "سبز"],
        ],
        [5.0, 3.5, 4.0, 3.0],
    )

    add_heading(doc, "Burndown Data")
    add_table(
        doc,
        ["روز", "ایده‌آل", "واقعی"],
        [[str(i), str(ideal), str(actual)] for i, ideal, actual in [
            (0, 42, 42), (1, 38, 40), (2, 34, 35), (3, 30, 31), (4, 26, 28),
            (5, 22, 24), (6, 18, 19), (7, 14, 15), (8, 10, 10), (9, 6, 6), (10, 0, 2)
        ]],
        [3.0, 5.0, 5.0],
    )

    add_heading(doc, "Time Tracking")
    add_table(
        doc,
        ["عضو", "تحلیل/برنامه‌ریزی", "پیاده‌سازی", "مستندات", "تست/ارائه", "جمع"],
        [
            ["محمدامین پورمند", "7h", "26h", "8h", "6h", "47h"],
            ["محمدرضا آرمان پور", "4h", "2h", "8h", "4h", "18h"],
            ["محدثه حاتمی کیا", "4h", "5h", "9h", "5h", "23h"],
            ["جمع کل", "15h", "33h", "25h", "15h", "88h"],
        ],
        [4.0, 2.8, 2.8, 2.8, 2.8, 2.5],
    )

    add_heading(doc, "RACI Matrix")
    add_table(
        doc,
        ["فعالیت", "محمدامین", "محمدرضا", "محدثه"],
        [
            ["تعریف scope", "A/R", "C", "C"],
            ["ارزش اجتماعی", "C", "A/R", "C"],
            ["مدل ML", "A/R", "C", "I"],
            ["API", "A/R", "I", "C"],
            ["UI", "C", "I", "A/R"],
            ["KPI و ریسک", "C", "A/R", "C"],
            ["QA و مستندات", "C", "C", "A/R"],
            ["ویدئو و تحویل", "A/R", "R", "R"],
        ],
        [5.8, 3.2, 3.2, 3.2],
    )

    add_heading(doc, "Work Management Board")
    add_table(
        doc,
        ["ستون", "هدف"],
        [
            ["Backlog", "همه taskهای محصول، docs و تحویل"],
            ["Selected for Sprint", "آیتم‌های sprint جاری"],
            ["In Progress", "کارهای در حال انجام"],
            ["Review/QA", "نیازمند بازبینی"],
            ["Done", "خروجی قابل مشاهده و ثبت‌شده"],
        ],
        [5.0, 10.0],
    )
    add_para(doc, "برای نمره کامل، همین ساختار باید در Trello یا Notion واقعی ساخته شود و taskها با assignee و time tracking وارد شوند.")

    add_heading(doc, "Knowledge Management")
    add_table(
        doc,
        ["صفحه", "مالک", "وضعیت"],
        [
            ["Project Home", "محمدامین", "آماده"],
            ["Problem & Stakeholders", "محمدرضا", "آماده"],
            ["Architecture", "محمدامین", "آماده"],
            ["API Documentation", "محمدامین", "آماده"],
            ["Risk Register", "محمدرضا", "آماده"],
            ["Meeting Notes", "محدثه", "آماده"],
            ["AI Usage Report", "محدثه", "آماده"],
        ],
        [6.0, 4.0, 4.0],
    )

    add_heading(doc, "Meeting Minutes خلاصه")
    add_table(
        doc,
        ["جلسه", "خروجی", "تصمیم"],
        [
            ["تعریف مسئله", "انتخاب تریاژ اورژانس", "decision-support، نه جایگزینی متخصص"],
            ["طراحی MVP", "API-first و mobile-first", "پشتیبانی از ورودی ناقص"],
            ["ارزیابی مدل", "انتخاب v6", "ثبت safety-first و balanced mode"],
            ["تحویل نهایی", "پوستر، ویدئو، گزارش و board", "تقسیم ارائه بین سه عضو"],
            ["بازخوانی rubric", "چک‌لیست مادر", "ساخت اسناد مستقل برای همه معیارها"],
        ],
        [4.5, 6.0, 6.0],
    )

    doc.save(OUT / "ITPM_Project_Management_Evidence_Package.docx")


def presentation_guide() -> None:
    doc = Document()
    style_doc(
        doc,
        "راهنمای ارائه، ویدئو و تحویل نهایی",
        "سناریوی ۱۰ دقیقه‌ای، demo، نقش اعضا و checklist روز تحویل",
    )

    add_heading(doc, "تقسیم زمان ویدئوی ۱۰ دقیقه‌ای")
    add_table(
        doc,
        ["زمان", "ارائه‌دهنده", "محتوا"],
        [
            ["0:00-0:45", "محمدامین", "معرفی پروژه و مسئله اورژانس"],
            ["0:45-1:45", "محمدرضا", "ارزش اجتماعی، ملی و مدیریت بحران"],
            ["1:45-3:00", "محمدامین", "معماری سیستم و MVP"],
            ["3:00-4:20", "محمدامین", "مدل v6، داده triage-time و leakage control"],
            ["4:20-5:20", "محمدرضا", "KPIها و trade-off safety-first"],
            ["5:20-7:00", "محدثه", "demo UI موبایل/PWA و سناریوهای test"],
            ["7:00-8:00", "محدثه", "QA، docs و ethics"],
            ["8:00-9:10", "محمدرضا", "چالش‌ها و درس‌آموخته‌ها"],
            ["9:10-10:00", "محمدامین", "roadmap و جمع‌بندی"],
        ],
        [3.0, 3.5, 9.5],
    )

    add_heading(doc, "Demo Checklist")
    add_bullets(
        doc,
        [
            "قبل از ضبط، API را اجرا کنید و /health را چک کنید.",
            "سناریوی بحرانی را با SpO2 پایین و shock index بالا نمایش دهید.",
            "سناریوی sparse را با ۴ آیتم نمایش دهید و data completeness را توضیح دهید.",
            "سناریوی cardiac را نمایش دهید و تفاوت model_probability با safety_flags را توضیح دهید.",
            "در viewport موبایل، action bar، تب‌ها و PWA install را نشان دهید.",
            "روی disclaimer تصمیم‌یار بودن تاکید کنید.",
            "اگر API هنگام ارائه مشکل داشت، از screenshot پوستر و خروجی‌های ثبت‌شده استفاده کنید.",
        ],
    )

    add_heading(doc, "جمله‌های کلیدی آماده")
    add_callout(
        doc,
        "ارزش انسانی",
        "در تریاژ، هزینه از دست دادن بیمار بحرانی از هزینه بررسی بیشتر چند بیمار غیر بحرانی بالاتر است.",
    )
    add_callout(
        doc,
        "Scope Pivot",
        "موضوع اولیه تغییر کرد چون مسئله تریاژ اورژانس با بحران، اثر اجتماعی و KPIهای کمی هم‌راستاتر بود.",
    )
    add_callout(
        doc,
        "اخلاق AI",
        "این سامانه پزشک نیست؛ فقط یک سیگنال کمکی قابل توضیح برای تصمیم بهتر است.",
    )

    add_heading(doc, "Checklist قبل از تحویل")
    add_table(
        doc,
        ["آیتم", "وضعیت پیشنهادی"],
        [
            ["پوستر A0 نهایی", "جایگزینی current-poster.png با نسخه نهایی"],
            ["گزارش Word/PDF", "استفاده از فایل گزارش نهایی رسمی"],
            ["بسته شواهد مدیریت پروژه", "ضمیمه یا لینک GitHub"],
            ["ویدئوی ۱۰ دقیقه‌ای", "حضور هر سه عضو"],
            ["Trello/Notion واقعی", "کپی taskها و time tracking"],
            ["Knowledge Base", "کپی ساختار docs"],
            ["GitHub", "آخرین commit push شده"],
        ],
        [8.0, 8.0],
    )

    doc.save(OUT / "ITPM_Presentation_and_Submission_Guide.docx")


def main() -> None:
    final_report()
    management_package()
    presentation_guide()
    print(f"Created official docs in {OUT}")


if __name__ == "__main__":
    main()
