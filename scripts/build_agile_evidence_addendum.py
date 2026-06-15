from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "deliverables" / "ITPM_Agile_Evidence_Addendum.docx"
FONT = "Tahoma"
TEAL = "087F8C"
NAVY = "0B2545"


def rtl(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))


def run_font(run, size=10.5, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:cs"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    rtl(p)
    r = p.add_run(text)
    run_font(r, bold=bold)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    rtl(p)
    for r in p.runs:
        run_font(r, size=16 if level == 1 else 12.5, bold=True, color=TEAL if level == 1 else NAVY)


def table(doc, headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    for i, header in enumerate(headers):
        p = tbl.rows[0].cells[i].paragraphs[0]
        rtl(p)
        r = p.add_run(header)
        run_font(r, 9.5, True, NAVY)
    for row in rows:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            rtl(p)
            r = p.add_run(str(value))
            run_font(r, 9)
    doc.add_paragraph()


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("پیوست تکمیلی شواهد Agile و پاسخ به بازخورد TA")
    run_font(r, 18, True, NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("سیستم هوشمند پشتیبان تصمیم‌گیری تریاژ اورژانس")
    run_font(r, 11, False, TEAL)

    heading(doc, "هدف سند")
    para(
        doc,
        "این پیوست برای پاسخ مستقیم به نقد TA تهیه شده است: پروژه نباید فقط یک پروژه فنی یا ML دیده شود، بلکه باید شواهد مدیریت پروژه، Sprint، task owner، ابزار مدیریت کار، مدیریت دانش و feedback loop داشته باشد.",
    )

    heading(doc, "ابزارهای مدیریت پروژه")
    table(
        doc,
        ["نیاز ارزیابی", "بستر پروژه", "مدرک"],
        [
            ["Backlog و Sprint", "GitHub Issues/Projects-ready", "docs/artifacts/github-issues-seed.csv"],
            ["Knowledge Management", "Repo Knowledge Base", "docs/knowledge-base/"],
            ["Decision Log", "Markdown docs", "docs/decision-log.md"],
            ["Stakeholder Feedback", "Feedback form in MVP", "frontend/index.html"],
            ["KPI/Burndown", "Agile dashboard", "docs/agile-dashboard.md"],
        ],
    )

    heading(doc, "Sprintها و Deliverableها")
    table(
        doc,
        ["Sprint", "هدف", "Deliverable", "وضعیت"],
        [
            ["Sprint 0", "تعریف مسئله و pivot", "Scope record و stakeholder register", "Done"],
            ["Sprint 1", "Baseline مدل", "مدل اولیه و کنترل leakage", "Done"],
            ["Sprint 2", "API MVP", "FastAPI و endpointهای اصلی", "Done"],
            ["Sprint 3", "مدل v6 و KPI", "metrics، model card و risk register", "Done"],
            ["Sprint 4", "Mobile/PWA و تحویل", "UI موبایل، PWA و Word deliverables", "Done"],
            ["Sprint 5", "PM evidence و feedback", "GitHub issue seed، Knowledge Base، feedback form", "In Progress"],
        ],
    )

    heading(doc, "نقش اعضای تیم")
    table(
        doc,
        ["عضو", "نقش", "بخش قابل ارائه"],
        [
            ["محمدامین پورمند", "Project Lead / ML & System Architect", "مدل، API، معماری، leakage و safety-first"],
            ["محمدرضا آرمان پور", "Project Control & Stakeholder/KPI Coordinator", "KPI، ریسک، burndown، ارزش اجتماعی و feedback"],
            ["محدثه حاتمی کیا", "UI/Documentation & QA Coordinator", "UI موبایل، QA، مستندات و سناریوهای تست"],
        ],
    )

    heading(doc, "برنامه بازخورد ذی‌نفع")
    para(doc, "در Sprint 5 فرم بازخورد داخل MVP اضافه شد. هدف تا ارائه نهایی دریافت حداقل 5 بازخورد واقعی است. بازخوردها علاوه بر fallback مرورگر، از طریق backend نیز در CSV سروری ذخیره و export می‌شوند.")
    table(
        doc,
        ["نوع ذی‌نفع", "هدف"],
        [
            ["دانشجوی پزشکی/پرستاری", "اعتبار اولیه زبان و خروجی بالینی"],
            ["آشنا با درمانگاه/اورژانس", "بررسی کاربردپذیری در محیط واقعی‌تر"],
            ["دانشجوی مهندسی", "بررسی UI و فهم خروجی"],
            ["کاربر عمومی", "بررسی شفاف بودن disclaimer"],
            ["آشنا با مدیریت بحران", "بررسی ارزش در شرایط ازدحام و بحران"],
        ],
    )
    table(
        doc,
        ["Endpoint", "هدف"],
        [
            ["POST /feedback", "ثبت بازخورد ذی‌نفع"],
            ["GET /feedback-summary", "نمایش تعداد بازخوردهای ثبت‌شده"],
            ["GET /feedback/export", "خروجی CSV برای گزارش نهایی"],
        ],
    )

    heading(doc, "جمله دفاعی آماده")
    para(
        doc,
        "ما Jira رسمی استفاده نکردیم چون برای مقیاس پروژه overhead زیادی داشت؛ اما artifactهای معادل Jira شامل backlog، sprint board، task assignment، time tracking، burndown، KPI tracking و knowledge base را در GitHub/repo docs پیاده‌سازی کردیم.",
        bold=True,
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
