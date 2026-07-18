from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
PATH = ROOT / "docs" / "deliverables" / "ITPM_Final_Announcement_Compliance_Matrix.docx"


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:bidi")) is None:
        p_pr.append(OxmlElement("w:bidi"))
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9)
    r_pr = run._r.get_or_add_rPr()
    r_pr.rFonts.set(qn("w:ascii"), "Arial")
    r_pr.rFonts.set(qn("w:hAnsi"), "Arial")
    r_pr.rFonts.set(qn("w:cs"), "Arial")
    if r_pr.find(qn("w:rtl")) is None:
        r_pr.append(OxmlElement("w:rtl"))


def update_row(table, label: str, values: dict[int, str]) -> None:
    for row in table.rows:
        if row.cells[0].text.strip() == label:
            for index, value in values.items():
                set_cell_text(row.cells[index], value)
            return
    raise RuntimeError(f"Compliance row not found: {label}")


def main() -> None:
    document = Document(PATH)
    tools_table = document.tables[2]
    update_row(
        tools_table,
        "Jira: وظیفه دقیق، Assignee و Agile",
        {
            1: "۹ Epic، ۵۱ Story/Task، Workflow، Owner، Sprint و DoD",
            2: "اجراشده و ممیزی‌شده",
            3: "دعوت حساب واقعی محدثه و محمدرضا برای Assignee شخصی",
        },
    )
    update_row(
        tools_table,
        "Jira: Time Tracking از کارهای انجام‌شده",
        {
            2: "ساختار و برآورد ثبت شده",
            3: "Worklog شخصی هر عضو پس از دعوت حساب واقعی تکمیل شود",
        },
    )
    update_row(
        tools_table,
        "Notion: ساختار اسناد",
        {
            1: "۸ صفحه محتوایی، ۵ دیتابیس و ۷۵ رکورد",
            2: "اجراشده و بازخوانی‌شده",
            3: "ثبت تصاویر احراز هویت‌شده و آزمون دسترسی مهمان",
        },
    )
    update_row(
        tools_table,
        "Notion: API و تصمیم فنی",
        {
            1: "API Docs، Architecture، Decision Log و Evidence",
            2: "اجراشده",
            3: "لینک commit نهایی پس از انتشار main به‌روزرسانی شود",
        },
    )
    update_row(
        tools_table,
        "Notion: صورت‌جلسه و دانش دوره",
        {
            2: "اجراشده با تاریخ‌های مرجع",
            3: "حضور فاقد شاهد همچنان قطعی گزارش نشود",
        },
    )

    general_table = document.tables[3]
    update_row(
        general_table,
        "گزارش Word راست‌چین",
        {1: "به‌روزرسانی‌شده و کنترل ساختاری شده؛ رندر LibreOffice در این محیط در دسترس نبود"},
    )
    update_row(
        general_table,
        "ZIP نهایی",
        {1: "بسته هسته آماده است؛ نسخه دارای شواهد پلتفرم پس از افزودن ۱۱ تصویر احراز هویت‌شده بازبسته‌بندی می‌شود"},
    )
    document.save(PATH)
    print(PATH)


if __name__ == "__main__":
    main()
